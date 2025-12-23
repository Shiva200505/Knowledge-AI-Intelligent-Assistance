"""
Document processing service for extracting text and metadata from uploaded files
"""
import os
import uuid
import asyncio
import aiofiles
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime

# Document processing libraries
import PyPDF2
from docx import Document
import pandas as pd
import json

from models.schemas import DocumentMetadata, DocumentChunk
from config.settings import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing and extracting content from documents"""
    
    def __init__(self, storage_path: str):
        self.storage_path = Path(storage_path)
        self.upload_path = Path(settings.upload_path)
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_text,
            'application/vnd.ms-excel': self._process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel
        }
        
    async def initialize(self):
        """Initialize storage directories"""
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.upload_path.mkdir(parents=True, exist_ok=True)
        logger.info("Document processor initialized")
        
    async def save_file(self, file, filename: str) -> str:
        """Save uploaded file to storage"""
        file_path = self.upload_path / filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
            
        logger.info(f"File saved: {filename}")
        return str(file_path)
        
    async def process_document(
        self, 
        file_path: str, 
        metadata: DocumentMetadata
    ) -> Dict[str, Any]:
        """Process document and extract structured content"""
        try:
            # Determine processor based on content type
            processor = self.supported_types.get(metadata.content_type)
            if not processor:
                raise ValueError(f"Unsupported content type: {metadata.content_type}")
                
            # Extract text content
            content_data = await processor(file_path)
            
            # Create document chunks
            chunks = await self._create_chunks(
                content_data['text'],
                metadata.id,
                content_data.get('structure', {})
            )
            
            # Extract metadata
            extracted_metadata = await self._extract_metadata(
                content_data,
                metadata
            )
            
            document_data = {
                'id': metadata.id,
                'metadata': extracted_metadata,
                'chunks': chunks,
                'full_text': content_data['text'],
                'structure': content_data.get('structure', {}),
                'images': content_data.get('images', []),
                'tables': content_data.get('tables', [])
            }
            
            # Move file to permanent storage
            permanent_path = self.storage_path / f"{metadata.id}_{metadata.filename}"
            await self._move_file(file_path, str(permanent_path))
            
            # Update metadata with actual permanent path
            document_data['final_path'] = str(permanent_path)
            
            logger.info(f"Document processed: {metadata.id}")
            return document_data
            
        except Exception as e:
            logger.error(f"Document processing failed for {metadata.id}: {e}")
            raise
            
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document - optimized for speed"""
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        
        def extract_page_text(page_num: int, page) -> tuple:
            """Extract text from a single page"""
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    return (page_num, page_text, {
                        'page_number': page_num,
                        'text_length': len(page_text),
                        'has_text': True
                    })
                return (page_num, None, None)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                return (page_num, None, None)
        
        text_content = []
        structure = {'pages': []}
        
        # Open PDF and get all pages first
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            # For small PDFs, process sequentially
            if total_pages <= 10:
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_num_res, page_text, page_struct = extract_page_text(page_num, page)
                    if page_text:
                        text_content.append(page_text)
                        structure['pages'].append(page_struct)
            else:
                # For larger PDFs, process in parallel (in thread pool)
                loop = asyncio.get_event_loop()
                with ThreadPoolExecutor(max_workers=4) as executor:
                    # Process pages in parallel
                    tasks = []
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        task = loop.run_in_executor(
                            executor,
                            extract_page_text,
                            page_num,
                            page
                        )
                        tasks.append(task)
                    
                    # Wait for all pages to complete
                    results = await asyncio.gather(*tasks)
                    
                    # Sort results by page number and collect text
                    results.sort(key=lambda x: x[0])
                    for page_num, page_text, page_struct in results:
                        if page_text:
                            text_content.append(page_text)
                            structure['pages'].append(page_struct)
            
            # Check if PDF is image-based (no extractable text)
            if len(text_content) == 0 and total_pages > 0:
                logger.warning(f"PDF appears to be image-based (no text extracted from {total_pages} pages)")
                text_content.append(f"[Image-based PDF]\n\nThis PDF contains {total_pages} pages but no extractable text. It may be a scanned document or contain only images. OCR processing would be required to extract text from this document.")
                    
        return {
            'text': '\n\n'.join(text_content),
            'structure': structure,
            'page_count': total_pages
        }
        
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        doc = Document(file_path)
        
        text_content = []
        structure = {'paragraphs': [], 'tables': []}
        
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                structure['paragraphs'].append({
                    'paragraph_number': i + 1,
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else None
                })
                
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(row_text)
            
            structure['tables'].append({
                'table_number': i + 1,
                'data': table_text
            })
            
            # Add table text to content
            table_content = '\n'.join(['\t'.join(row) for row in table_text])
            text_content.append(f"[Table {i + 1}]\n{table_content}")
            
        return {
            'text': '\n\n'.join(text_content),
            'structure': structure
        }
        
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text document"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            text = await file.read()
            
        # Simple structure detection
        lines = text.split('\n')
        structure = {
            'lines': len(lines),
            'paragraphs': len([line for line in lines if line.strip()])
        }
        
        return {
            'text': text,
            'structure': structure
        }
        
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel spreadsheet"""
        df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
        
        text_content = []
        structure = {'sheets': []}
        
        for sheet_name, sheet_df in df.items():
            # Convert sheet to text
            sheet_text = f"[Sheet: {sheet_name}]\n"
            sheet_text += sheet_df.to_string(index=False)
            text_content.append(sheet_text)
            
            structure['sheets'].append({
                'name': sheet_name,
                'rows': len(sheet_df),
                'columns': len(sheet_df.columns),
                'column_names': list(sheet_df.columns)
            })
            
        return {
            'text': '\n\n'.join(text_content),
            'structure': structure
        }
        
    async def _create_chunks(
        self, 
        text: str, 
        document_id: str,
        structure: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create semantic chunks from document text - optimized"""
        chunks = []
        
        # Use smaller chunk size for faster processing (500 words instead of 1000)
        max_chunk_size = min(500, settings.context_window_size)
        
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Quick word count estimation (faster than split().len())
        def quick_word_count(text: str) -> int:
            return text.count(' ') + text.count('\n') + 1
        
        current_chunk = ""
        current_size = 0
        chunk_index = 0
        
        for i, paragraph in enumerate(paragraphs):
            paragraph_size = quick_word_count(paragraph)
            
            # Check if adding this paragraph would exceed chunk size
            if current_size + paragraph_size > max_chunk_size and current_chunk:
                # Create chunk
                chunk = DocumentChunk(
                    id=f"{document_id}_chunk_{chunk_index}",
                    document_id=document_id,
                    content=current_chunk.strip(),
                    page_number=self._estimate_page_number(chunk_index, structure),
                    paragraph_number=i,
                    chunk_index=chunk_index,
                    metadata={
                        'word_count': current_size,
                        'paragraph_count': current_chunk.count('\n\n') + 1
                    }
                )
                chunks.append(chunk)
                
                # Reset for next chunk
                current_chunk = paragraph
                current_size = paragraph_size
                chunk_index += 1
            else:
                # Add to current chunk
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
                current_size += paragraph_size
                
        # Add final chunk
        if current_chunk:
            chunk = DocumentChunk(
                id=f"{document_id}_chunk_{chunk_index}",
                document_id=document_id,
                content=current_chunk.strip(),
                page_number=self._estimate_page_number(chunk_index, structure),
                paragraph_number=len(paragraphs),
                chunk_index=chunk_index,
                metadata={
                    'word_count': current_size,
                    'paragraph_count': current_chunk.count('\n\n') + 1
                }
            )
            chunks.append(chunk)
            
        logger.info(f"Created {len(chunks)} chunks for document {document_id}")
        return chunks
        
    def _estimate_page_number(self, chunk_index: int, structure: Dict[str, Any]) -> int:
        """Estimate page number for chunk based on document structure"""
        if 'pages' in structure:
            total_pages = len(structure['pages'])
            if total_pages > 0:
                # Rough estimate based on chunk position
                estimated_page = min(chunk_index + 1, total_pages)
                return estimated_page
        
        # Default to page 1 if no structure info
        return 1
        
    async def _extract_metadata(
        self, 
        content_data: Dict[str, Any],
        original_metadata: DocumentMetadata
    ) -> Dict[str, Any]:
        """Extract enhanced metadata from document content"""
        text = content_data['text']
        
        metadata = {
            'word_count': len(text.split()),
            'character_count': len(text),
            'page_count': content_data.get('page_count', 1),
            'language': 'en',  # Could use language detection
            'structure': content_data.get('structure', {}),
            'processed_at': datetime.now().isoformat(),
            'file_size': os.path.getsize(original_metadata.file_path) if os.path.exists(original_metadata.file_path) else 0
        }
        
        # Merge with original metadata
        metadata.update(original_metadata.metadata)
        
        return metadata
        
    async def _move_file(self, source: str, destination: str):
        """Move file from temporary to permanent storage"""
        source_path = Path(source)
        destination_path = Path(destination)
        
        if source_path.exists():
            destination_path.parent.mkdir(parents=True, exist_ok=True)
            source_path.rename(destination_path)
            
    async def delete_document(self, document_id: str):
        """Delete document file from storage"""
        try:
            # Find and delete document files
            for file_path in self.storage_path.glob(f"{document_id}_*"):
                file_path.unlink()
                
            logger.info(f"Document files deleted: {document_id}")
            
        except Exception as e:
            logger.error(f"Failed to delete document files for {document_id}: {e}")
            raise